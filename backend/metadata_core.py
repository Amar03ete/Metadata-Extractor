"""
Core metadata extraction engine.
Extracts OS-level and document-level metadata from various file types.
"""
import os
import platform
import hashlib
import socket
from datetime import datetime, timezone, timedelta
from pathlib import Path
from typing import Dict, Any, Optional
import psutil

# Optional Unix-only utilities (pwd/grp) - guard their usage on Windows
try:
    import pwd
    import grp
    HAS_UNIX_UTILS = True
except Exception:
    pwd = None
    grp = None
    HAS_UNIX_UTILS = False
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# PDF libraries (optional)
try:
    import PyPDF2
    HAS_PDF = True
except Exception:
    PyPDF2 = None
    HAS_PDF = False

try:
    import pikepdf
except Exception:
    pikepdf = None

try:
    from openpyxl import load_workbook
    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False

try:
    from pptx import Presentation
    HAS_PPTX = True
except ImportError:
    HAS_PPTX = False


def extract_filesystem_metadata(file_path: str) -> Dict[str, Any]:
    """Extract OS-level filesystem metadata."""
    stat_info = os.stat(file_path)
    path_obj = Path(file_path)
    
    # Convert timestamps with timezone awareness
    if platform.system() == 'Windows':
        created_dt = datetime.fromtimestamp(stat_info.st_ctime)
    else:
        try:
            created_dt = datetime.fromtimestamp(stat_info.st_birthtime)
        except AttributeError:
            created_dt = datetime.fromtimestamp(stat_info.st_ctime)
    
    modified_dt = datetime.fromtimestamp(stat_info.st_mtime)
    accessed_dt = datetime.fromtimestamp(stat_info.st_atime)
    
    # Add timezone info to timestamps
    local_tz = datetime.now(timezone.utc).astimezone().tzinfo
    if local_tz:
        created_dt = created_dt.replace(tzinfo=local_tz)
        modified_dt = modified_dt.replace(tzinfo=local_tz)
        accessed_dt = accessed_dt.replace(tzinfo=local_tz)
    
    tz_offset = local_tz.utcoffset(datetime.now()) if local_tz else None
    tz_name = local_tz.tzname(datetime.now()) if local_tz else "Unknown"
    
    # Calculate file hashes
    file_hashes = calculate_file_hashes(file_path)
    
    # Get owner/group names when available (Unix-like systems)
    owner_name = "Unknown"
    group_name = "Unknown"
    if HAS_UNIX_UTILS:
        try:
            owner_name = pwd.getpwuid(stat_info.st_uid).pw_name
            group_name = grp.getgrgid(stat_info.st_gid).gr_name
        except Exception:
            owner_name = str(stat_info.st_uid)
            group_name = str(stat_info.st_gid)
    else:
        # On Windows or when pwd/grp not available, keep numeric IDs as fallback
        owner_name = str(stat_info.st_uid)
        group_name = str(stat_info.st_gid)
    
    # Determine MIME type
    mime_type = "application/octet-stream"
    ext = path_obj.suffix.lower()
    mime_types = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.doc': 'application/msword',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.xls': 'application/vnd.ms-excel',
        '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
        '.ppt': 'application/vnd.ms-powerpoint',
    }
    mime_type = mime_types.get(ext, mime_type)
    
    metadata = {
        # File Identity
        "filename": path_obj.name,
        "file_path": str(path_obj.absolute()),
        "file_extension": ext,
        "mime_type": mime_type,
        
        # File Size
        "size_bytes": stat_info.st_size,
        "size_formatted": format_file_size(stat_info.st_size),
        
        # Unique Identifiers
        "file_hash_md5": file_hashes.get('md5'),
        "file_hash_sha1": file_hashes.get('sha1'),
        "file_hash_sha256": file_hashes.get('sha256'),
        "inode": stat_info.st_ino,
        "hard_links": stat_info.st_nlink,
        
        # Timestamps with Timezone
        "fs_created": created_dt.isoformat() if created_dt else None,
        "fs_created_timestamp": stat_info.st_ctime,
        "fs_modified": modified_dt.isoformat() if modified_dt else None,
        "fs_modified_timestamp": stat_info.st_mtime,
        "fs_accessed": accessed_dt.isoformat() if accessed_dt else None,
        "fs_accessed_timestamp": stat_info.st_atime,
        
        # Timezone Information
        "timezone_name": tz_name,
        "timezone_offset": str(tz_offset),
        "timezone_offset_seconds": int(tz_offset.total_seconds()) if tz_offset else None,
        
        # Storage and Access
        "file_permissions": oct(stat_info.st_mode)[-3:],
        "file_permissions_octal": oct(stat_info.st_mode),
        "owner_uid": stat_info.st_uid,
        "owner_name": owner_name,
        "group_gid": stat_info.st_gid,
        "group_name": group_name,
        
        # OS/Platform Information
        "system": platform.system(),
        "system_version": platform.release(),
        "system_architecture": platform.machine(),
    }
    
    return metadata


def format_file_size(bytes_size: int) -> str:
    """Format bytes to human readable size."""
    for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
        if bytes_size < 1024.0:
            return f"{bytes_size:.2f} {unit}"
        bytes_size /= 1024.0
    return f"{bytes_size:.2f} PB"


def calculate_file_hashes(file_path: str) -> Dict[str, str]:
    """Calculate MD5, SHA-1, and SHA-256 hashes of file."""
    hashes = {}
    try:
        md5 = hashlib.md5()
        sha1 = hashlib.sha1()
        sha256 = hashlib.sha256()
        
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(8192), b''):
                md5.update(chunk)
                sha1.update(chunk)
                sha256.update(chunk)
        
        hashes['md5'] = md5.hexdigest()
        hashes['sha1'] = sha1.hexdigest()
        hashes['sha256'] = sha256.hexdigest()
    except Exception as e:
        hashes['error'] = str(e)
    
    return hashes


def get_os_hardware_info() -> Dict[str, Any]:
    """Extract comprehensive OS and hardware information."""
    import sys
    info = {
        "os_system": platform.system(),
        "os_release": platform.release(),
        "os_version": platform.version(),
        "os_machine": platform.machine(),
        "os_platform": sys.platform,
        "hostname": socket.gethostname(),
        "processor": platform.processor(),
        "cpu_count_physical": psutil.cpu_count(logical=False),
        "cpu_count_logical": psutil.cpu_count(logical=True),
        "cpu_percent": psutil.cpu_percent(interval=0.1),
        "memory_total": format_file_size(psutil.virtual_memory().total),
        "memory_available": format_file_size(psutil.virtual_memory().available),
        "memory_percent": psutil.virtual_memory().percent,
        "python_version": platform.python_version(),
        "python_implementation": platform.python_implementation(),
    }
    
    return info


def get_timezone_info() -> Dict[str, Any]:
    """Extract detailed timezone information."""
    try:
        import time
        import sys
        
        # Get local timezone
        local_tz = datetime.now(timezone.utc).astimezone().tzinfo
        local_offset = local_tz.utcoffset(datetime.now()) if local_tz else None
        local_name = local_tz.tzname(datetime.now()) if local_tz else "Unknown"
        
        # Get system timezone
        try:
            tz_abbr = time.tzname
            is_dst = time.daylight
        except:
            tz_abbr = ("Unknown", "Unknown")
            is_dst = False
        
        return {
            "local_timezone": local_name,
            "utc_offset": str(local_offset) if local_offset else None,
            "utc_offset_seconds": int(local_offset.total_seconds()) if local_offset else None,
            "is_daylight_saving": is_dst,
            "timezone_names": {
                "standard": tz_abbr[0] if len(tz_abbr) > 0 else "Unknown",
                "daylight": tz_abbr[1] if len(tz_abbr) > 1 else "Unknown"
            }
        }
    except Exception as e:
        return {"error": str(e)}


def extract_pdf_metadata(file_path: str) -> Dict[str, Any]:
    """Extract metadata from PDF files."""
    metadata = {}
    
    if not HAS_PDF and not pikepdf:
        return {"pdf_error": "PDF libraries not installed"}
    
    # Try PyPDF2 first if available
    if PyPDF2:
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                if getattr(pdf_reader, 'metadata', None):
                    pdf_meta = pdf_reader.metadata
                    metadata["pdf_title"] = pdf_meta.get('/Title', '')
                    metadata["pdf_author"] = pdf_meta.get('/Author', '')
                    metadata["pdf_subject"] = pdf_meta.get('/Subject', '')
                    metadata["pdf_creator"] = pdf_meta.get('/Creator', '')
                    metadata["pdf_producer"] = pdf_meta.get('/Producer', '')
                    metadata["pdf_creationdate"] = pdf_meta.get('/CreationDate', '')
                    metadata["pdf_moddate"] = pdf_meta.get('/ModDate', '')
                    metadata["pdf_keywords"] = pdf_meta.get('/Keywords', '')
                    metadata["pdf_trapped"] = pdf_meta.get('/Trapped', '')
                try:
                    metadata["pdf_pages"] = len(pdf_reader.pages)
                except Exception:
                    pass
        except Exception as e:
            metadata["pdf_error_pypdf2"] = str(e)

    # Try pikepdf for more comprehensive metadata if available
    if pikepdf:
        try:
            with pikepdf.Pdf.open(file_path) as pdf:
                metadata["pdf_encryption"] = "Yes" if pdf.is_encrypted else "No"
                metadata["pdf_version"] = str(getattr(pdf, 'pdf_version', ''))
                metadata["pdf_is_linearized"] = getattr(pdf, 'is_linearized', False)

                if getattr(pdf, 'docinfo', None):
                    for key, value in pdf.docinfo.items():
                        try:
                            metadata[f"pdf_{key.lower()}"] = str(value)
                        except Exception:
                            pass
        except Exception as e:
            metadata["pdf_error_pikepdf"] = str(e)
    
    return metadata


def extract_docx_metadata(file_path: str) -> Dict[str, Any]:
    """Extract metadata from DOCX files."""
    metadata = {}
    
    if not HAS_DOCX:
        return {"docx_error": "python-docx not installed"}
    
    try:
        doc = Document(file_path)
        core_props = doc.core_properties
        
        metadata["docx_title"] = core_props.title or ''
        metadata["docx_author"] = core_props.author or ''
        metadata["docx_subject"] = core_props.subject or ''
        metadata["docx_created"] = core_props.created.isoformat() if core_props.created else None
        metadata["docx_created_timestamp"] = core_props.created.timestamp() if core_props.created else None
        metadata["docx_modified"] = core_props.modified.isoformat() if core_props.modified else None
        metadata["docx_modified_timestamp"] = core_props.modified.timestamp() if core_props.modified else None
        metadata["docx_last_modified_by"] = core_props.last_modified_by or ''
        metadata["docx_revision"] = core_props.revision or ''
        metadata["docx_category"] = core_props.category or ''
        metadata["docx_comments"] = core_props.comments or ''
        metadata["docx_keywords"] = core_props.keywords or ''
        metadata["docx_language"] = core_props.language or ''
        metadata["docx_company"] = core_props.company or ''
        metadata["docx_manager"] = core_props.manager or ''
        
        # Count paragraphs and images
        try:
            metadata["docx_paragraphs"] = len(doc.paragraphs)
            metadata["docx_tables"] = len(doc.tables)
        except:
            pass
    except Exception as e:
        metadata["docx_error"] = str(e)
    
    return metadata


def extract_xlsx_metadata(file_path: str) -> Dict[str, Any]:
    """Extract metadata from XLSX files."""
    metadata = {}
    
    if not HAS_XLSX:
        return {"xlsx_error": "openpyxl not installed"}
    
    try:
        wb = load_workbook(file_path)
        props = wb.properties
        
        metadata["xlsx_title"] = props.title or ''
        metadata["xlsx_author"] = props.creator or ''
        metadata["xlsx_subject"] = props.subject or ''
        metadata["xlsx_created"] = props.created.isoformat() if props.created else None
        metadata["xlsx_created_timestamp"] = props.created.timestamp() if props.created else None
        metadata["xlsx_modified"] = props.modified.isoformat() if props.modified else None
        metadata["xlsx_modified_timestamp"] = props.modified.timestamp() if props.modified else None
        metadata["xlsx_last_modified_by"] = props.lastModifiedBy or ''
        metadata["xlsx_keywords"] = props.keywords or ''
        metadata["xlsx_category"] = props.category or ''
        metadata["xlsx_comments"] = props.description or ''
        metadata["xlsx_company"] = props.company or ''
        metadata["xlsx_manager"] = props.manager or ''
        
        # Count sheets
        try:
            metadata["xlsx_sheets"] = len(wb.sheetnames)
            metadata["xlsx_sheet_names"] = wb.sheetnames
        except:
            pass
    except Exception as e:
        metadata["xlsx_error"] = str(e)
    
    return metadata


def extract_pptx_metadata(file_path: str) -> Dict[str, Any]:
    """Extract metadata from PPTX files."""
    metadata = {}
    
    if not HAS_PPTX:
        return {"pptx_error": "python-pptx not installed"}
    
    try:
        prs = Presentation(file_path)
        core_props = prs.core_properties
        
        metadata["pptx_title"] = core_props.title or ''
        metadata["pptx_author"] = core_props.author or ''
        metadata["pptx_subject"] = core_props.subject or ''
        metadata["pptx_created"] = core_props.created.isoformat() if core_props.created else None
        metadata["pptx_created_timestamp"] = core_props.created.timestamp() if core_props.created else None
        metadata["pptx_modified"] = core_props.modified.isoformat() if core_props.modified else None
        metadata["pptx_modified_timestamp"] = core_props.modified.timestamp() if core_props.modified else None
        metadata["pptx_last_modified_by"] = core_props.last_modified_by or ''
        metadata["pptx_revision"] = core_props.revision or ''
        metadata["pptx_category"] = core_props.category or ''
        metadata["pptx_comments"] = core_props.comments or ''
        metadata["pptx_keywords"] = core_props.keywords or ''
        metadata["pptx_company"] = core_props.company or ''
        
        # Count slides
        try:
            metadata["pptx_slides"] = len(prs.slides)
        except:
            pass
    except Exception as e:
        metadata["pptx_error"] = str(e)
    
    return metadata


def extract_all_metadata(file_path: str) -> Dict[str, Any]:
    """
    Main function to extract all available metadata from a file.
    Returns a comprehensive dictionary with OS, hardware, and document-level metadata.
    """
    all_metadata = {}
    
    # Extract OS and hardware information
    all_metadata['os_hardware_metadata'] = get_os_hardware_info()
    
    # Extract timezone information
    all_metadata['timezone_metadata'] = get_timezone_info()
    
    # Always extract filesystem metadata
    all_metadata.update(extract_filesystem_metadata(file_path))
    
    # Extract document-specific metadata based on file extension
    file_ext = Path(file_path).suffix.lower()
    
    if file_ext == '.pdf':
        all_metadata.update(extract_pdf_metadata(file_path))
    elif file_ext in ['.docx', '.doc']:
        if file_ext == '.docx':
            all_metadata.update(extract_docx_metadata(file_path))
    elif file_ext in ['.xlsx', '.xls']:
        if file_ext == '.xlsx':
            all_metadata.update(extract_xlsx_metadata(file_path))
    elif file_ext in ['.pptx', '.ppt']:
        if file_ext == '.pptx':
            all_metadata.update(extract_pptx_metadata(file_path))
    
    return all_metadata

